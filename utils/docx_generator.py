from docx import Document
from docx.shared import Inches
import io
import os

def generate_docx(data, image_path=None):
    doc = Document()
    doc.add_heading("空間設計導覽文案", 0)

    # 設計理念
    doc.add_heading("設計理念總述", level=1)
    doc.add_paragraph(data.get("concept", ""))

    # 空間總覽＋插圖
    doc.add_heading("空間總覽與動線說明", level=1)
    doc.add_paragraph("以下為本案動線說明與原始平面圖：")
    if image_path and os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(5.5))

    # 空間逐區
    doc.add_heading("空間導覽文案", level=1)
    sections = data.get("sections", [])
    for section in sections:
        doc.add_heading(section.get("room", "未命名空間"), level=2)
        doc.add_paragraph(f"坪數：{section.get('area', '')}")
        doc.add_paragraph(f"功能：{section.get('function', '')}")
        doc.add_paragraph("傢俱配置：" + "、".join(section.get("furniture", [])))
        doc.add_paragraph(f"色彩搭配：{section.get('color', '')}")
        doc.add_paragraph(f"設計重點：{section.get('design_note', '')}")
        doc.add_paragraph(f"情感描述：{section.get('emotion', '')}")

    # 結語
    doc.add_heading("結語", level=1)
    doc.add_paragraph(data.get("conclusion", ""))

    # 若有品牌頁、屋主故事可擴充於此
    # if data.get("brand_info"):
    #     doc.add_heading("品牌故事", level=1)
    #     doc.add_paragraph(data["brand_info"])
    # if data.get("owner_story"):
    #     doc.add_heading("屋主故事", level=1)
    #     doc.add_paragraph(data["owner_story"])

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
